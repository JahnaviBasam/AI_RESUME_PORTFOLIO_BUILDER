# pyrefly: ignore [missing-import]
from docx import Document
import io

def generate_docx(text_content):
    """
    Generates a DOCX file from markdown/text content using python-docx.
    Returns a BytesIO object containing the DOCX.
    """
    doc = Document()
    
    for line in text_content.split('\n'):
        if line.strip():
            if line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('-') or line.startswith('*'):
                doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line)
                
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
